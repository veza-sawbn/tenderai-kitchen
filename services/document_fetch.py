from __future__ import annotations

import io
import os
import re
from datetime import datetime, timezone
from typing import Any, Dict, List, Optional

import requests
from pypdf import PdfReader
from sqlalchemy import desc, select

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

from models import TenderCache, TenderDocumentCache

DOC_FETCH_TIMEOUT = int(os.getenv("DOC_FETCH_TIMEOUT", "90"))
DOC_FETCH_RETRIES = int(os.getenv("DOC_FETCH_RETRIES", "3"))
DOC_FETCH_RETRY_SLEEP = float(os.getenv("DOC_FETCH_RETRY_SLEEP", "3"))


def utcnow():
    return datetime.now(timezone.utc)


def _safe_str(value: Any) -> Optional[str]:
    if value is None:
        return None
    text = str(value).strip()
    return text or None


def _filename_from_headers(response: requests.Response, fallback_url: str) -> Optional[str]:
    cd = response.headers.get("Content-Disposition") or ""
    match = re.search(r'filename\*?=(?:UTF-8\'\')?"?([^";]+)"?', cd, re.I)
    if match:
        return match.group(1).strip()
    if fallback_url:
        return fallback_url.split("/")[-1][:255] or None
    return None


def _looks_like_direct_document_url(url: str) -> bool:
    if not url:
        return False
    lower = url.lower()
    return any(lower.endswith(ext) for ext in [".pdf", ".docx", ".doc"]) or "download?" in lower or "blobname=" in lower


def _extract_pdf_text_from_bytes(binary: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(binary))
        parts = []
        for page in reader.pages:
            try:
                parts.append(page.extract_text() or "")
            except Exception:
                continue
        return "\n".join(parts).strip()
    except Exception:
        return ""


def _extract_docx_text_from_bytes(binary: bytes) -> str:
    if DocxDocument is None:
        return ""
    try:
        doc = DocxDocument(io.BytesIO(binary))
        parts: List[str] = []

        for p in doc.paragraphs:
            text = (p.text or "").strip()
            if text:
                parts.append(text)

        for table in doc.tables:
            for row in table.rows:
                row_parts = []
                for cell in row.cells:
                    cell_text = (cell.text or "").strip()
                    if cell_text:
                        row_parts.append(cell_text)
                if row_parts:
                    parts.append(" | ".join(row_parts))

        return "\n".join(parts).strip()
    except Exception:
        return ""


def latest_document_cache(session, tender_id: int) -> Optional[TenderDocumentCache]:
    return session.execute(
        select(TenderDocumentCache)
        .where(TenderDocumentCache.tender_id == tender_id)
        .order_by(desc(TenderDocumentCache.fetched_at), desc(TenderDocumentCache.id))
        .limit(1)
    ).scalars().first()


def _candidate_url(tender: TenderCache) -> Optional[str]:
    if tender.document_url:
        return tender.document_url
    if tender.source_url and _looks_like_direct_document_url(tender.source_url):
        return tender.source_url
    return None


def fetch_and_cache_document(
    session,
    tender: TenderCache,
    *,
    force_retry_failed: bool = False,
) -> Dict[str, Any]:
    url = _candidate_url(tender)
    if not url:
        return {
            "ok": False,
            "tender_id": tender.id,
            "status": "no_document_url",
            "error": "No direct tender document URL is available.",
        }

    existing = session.execute(
        select(TenderDocumentCache)
        .where(TenderDocumentCache.tender_id == tender.id, TenderDocumentCache.document_url == url)
        .order_by(desc(TenderDocumentCache.fetched_at), desc(TenderDocumentCache.id))
        .limit(1)
    ).scalars().first()

    if existing and existing.fetch_status in {"fetched", "fetched_unparsed", "unsupported_legacy_doc"} and not force_retry_failed:
        return {
            "ok": True,
            "tender_id": tender.id,
            "status": "skipped_existing",
            "fetch_status": existing.fetch_status,
        }

    if existing is None:
        existing = TenderDocumentCache(tender_id=tender.id, document_url=url)
        session.add(existing)
        session.flush()

    last_exc: Optional[Exception] = None
    response: Optional[requests.Response] = None
    for attempt in range(1, DOC_FETCH_RETRIES + 1):
        try:
            response = requests.get(url, timeout=DOC_FETCH_TIMEOUT, allow_redirects=True)
            response.raise_for_status()
            break
        except Exception as exc:
            last_exc = exc
            if attempt < DOC_FETCH_RETRIES:
                import time
                time.sleep(DOC_FETCH_RETRY_SLEEP * attempt)
                continue
            response = None

    if response is None:
        existing.fetch_status = "fetch_failed"
        existing.error_message = str(last_exc) if last_exc else "Unknown fetch failure."
        existing.fetched_at = utcnow()
        return {
            "ok": False,
            "tender_id": tender.id,
            "status": "fetch_failed",
            "error": existing.error_message,
        }

    binary = response.content or b""
    content_type = (response.headers.get("Content-Type") or "").lower()
    filename = _filename_from_headers(response, url)
    lower_name = (filename or "").lower()
    lower_url = (url or "").lower()

    is_pdf = "pdf" in content_type or lower_name.endswith(".pdf") or lower_url.endswith(".pdf")
    is_docx = "wordprocessingml" in content_type or lower_name.endswith(".docx") or lower_url.endswith(".docx")
    is_doc = (
        "msword" in content_type
        or lower_name.endswith(".doc")
        or (lower_url.endswith(".doc") and not lower_url.endswith(".docx"))
    )

    extracted_text = ""
    fetch_status = "fetched_unparsed"
    error_message = None

    if is_pdf:
        extracted_text = _extract_pdf_text_from_bytes(binary)
        fetch_status = "fetched" if extracted_text else "fetched_unparsed"
        if not extracted_text:
            error_message = "PDF fetched but no readable text was extracted."
    elif is_docx:
        extracted_text = _extract_docx_text_from_bytes(binary)
        fetch_status = "fetched" if extracted_text else "fetched_unparsed"
        if not extracted_text:
            error_message = "DOCX fetched but no readable text was extracted."
    elif is_doc:
        fetch_status = "unsupported_legacy_doc"
        error_message = "Legacy .doc file fetched but parsing is not supported in this pipeline."
    else:
        fetch_status = "fetched_unparsed"
        error_message = f"Unsupported or unknown document type: {content_type or 'unknown'}"

    existing.filename = filename
    existing.content_type = content_type[:100] if content_type else None
    existing.binary_content = binary
    existing.extracted_text = extracted_text or ""
    existing.fetch_status = fetch_status
    existing.error_message = error_message
    existing.fetched_at = utcnow()

    return {
        "ok": True,
        "tender_id": tender.id,
        "status": "fetched",
        "fetch_status": fetch_status,
        "parsed": bool(extracted_text),
        "filename": filename,
    }


def fetch_documents_for_live_tenders(
    session,
    *,
    limit: int = 10,
    force_retry_failed: bool = False,
) -> Dict[str, Any]:
    limit = max(1, min(int(limit), 100))

    tenders = session.execute(
        select(TenderCache)
        .where(
            TenderCache.is_live.is_(True),
            TenderCache.document_url.is_not(None),
        )
        .order_by(TenderCache.closing_date.asc().nulls_last(), desc(TenderCache.updated_at))
        .limit(limit)
    ).scalars().all()

    results = {
        "ok": True,
        "limit": limit,
        "candidates": len(tenders),
        "processed": 0,
        "fetched": 0,
        "parsed": 0,
        "skipped_existing": 0,
        "unsupported_legacy_doc": 0,
        "fetch_failed": 0,
        "items": [],
    }

    for tender in tenders:
        result = fetch_and_cache_document(session, tender, force_retry_failed=force_retry_failed)
        results["processed"] += 1
        results["items"].append(result)

        status = result.get("status")
        fetch_status = result.get("fetch_status")

        if status == "skipped_existing":
            results["skipped_existing"] += 1
        elif status == "fetch_failed":
            results["fetch_failed"] += 1
        elif status == "fetched":
            results["fetched"] += 1
            if result.get("parsed"):
                results["parsed"] += 1
            if fetch_status == "unsupported_legacy_doc":
                results["unsupported_legacy_doc"] += 1

    return results

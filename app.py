import io
import json
import os
import re
import tempfile
from datetime import date, datetime, timezone
from typing import Any, Dict, List, Optional

try:
    from dotenv import load_dotenv
except ImportError:
    def load_dotenv(*args, **kwargs):
        return False

import requests
from flask import (
    Flask,
    abort,
    flash,
    jsonify,
    redirect,
    render_template,
    request,
    send_file,
    url_for,
)
from pypdf import PdfReader
from sqlalchemy import desc, func, or_, select

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    from openai import OpenAI
except ImportError:
    OpenAI = None

from database import get_db_session, init_db
from models import AnalysisJob, IngestRun, Profile, ProfileIssue, TenderCache, TenderDocumentCache
from services.etenders_ingest import ingest_tenders


load_dotenv()


def utcnow():
    return datetime.now(timezone.utc)


app = Flask(__name__)
app.config["SECRET_KEY"] = os.getenv("FLASK_SECRET_KEY", "tenderai-dev-fallback-secret")
app.config["MAX_CONTENT_LENGTH"] = 25 * 1024 * 1024

init_db()

OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-4.1-mini")


def get_openai_client():
    api_key = (os.getenv("OPENAI_API_KEY") or "").strip()
    if not api_key or OpenAI is None:
        return None
    try:
        return OpenAI(api_key=api_key)
    except Exception:
        return None


def json_from_text(value: str) -> dict:
    if not value:
        return {}
    text = value.strip()
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?\s*", "", text)
        text = re.sub(r"\s*```$", "", text)
    try:
        return json.loads(text)
    except Exception:
        return {}


def safe_loads(value: Any, fallback=None):
    if fallback is None:
        fallback = {}
    if not value:
        return fallback
    if isinstance(value, dict):
        return value
    try:
        return json.loads(value)
    except Exception:
        return fallback


def admin_allowed() -> bool:
    token = (os.getenv("ADMIN_TOKEN") or "").strip()
    if not token:
        return True
    supplied = (request.headers.get("X-Admin-Token") or request.args.get("token") or "").strip()
    return supplied == token


def get_active_profile(session) -> Optional[Profile]:
    return session.execute(
        select(Profile)
        .where(Profile.is_active.is_(True))
        .order_by(desc(Profile.updated_at))
        .limit(1)
    ).scalars().first()


def normalize_keywords(text: str) -> List[str]:
    if not text:
        return []
    raw = text.replace("\n", ",").replace(";", ",").replace("|", ",").split(",")
    cleaned = []
    seen = set()
    for item in raw:
        value = " ".join(item.strip().split())
        if len(value) >= 3:
            key = value.lower()
            if key not in seen:
                seen.add(key)
                cleaned.append(value)
    return cleaned[:30]


def tokenize(text: str) -> List[str]:
    if not text:
        return []
    words = re.findall(r"[a-zA-Z0-9]+", text.lower())
    stop = {
        "the", "and", "for", "with", "from", "that", "this", "into", "your", "their",
        "supply", "supplies", "service", "services", "tender", "bid", "request",
        "proposal", "rfq", "rfp", "of", "to", "in", "on", "at", "by", "or", "an", "a"
    }
    out = []
    seen = set()
    for w in words:
        if len(w) < 4 or w in stop:
            continue
        if w not in seen:
            seen.add(w)
            out.append(w)
    return out


def extract_pdf_text(file_or_path) -> str:
    temp_path = None
    try:
        if isinstance(file_or_path, str):
            path = file_or_path
        else:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                file_or_path.save(tmp.name)
                path = tmp.name
                temp_path = tmp.name

        reader = PdfReader(path)
        parts = []
        for page in reader.pages:
            try:
                parts.append(page.extract_text() or "")
            except Exception:
                continue
        return "\n".join(parts).strip()
    finally:
        if temp_path:
            try:
                os.unlink(temp_path)
            except Exception:
                pass


def extract_docx_text(path: str) -> str:
    if DocxDocument is None:
        return ""
    try:
        doc = DocxDocument(path)
        parts = []

        for p in doc.paragraphs:
            text = (p.text or "").strip()
            if text:
                parts.append(text)

        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = (cell.text or "").strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    parts.append(" | ".join(row_text))

        return "\n".join(parts).strip()
    except Exception:
        return ""


def parse_profile_with_openai(text: str, filename: str | None = None) -> dict:
    client = get_openai_client()
    if not client or not text.strip():
        return {}

    prompt = f"""
You are TenderAI. Extract structured supplier profile information.
Return JSON only.

Schema:
{{
  "company_name": "string or null",
  "industry": "string or null",
  "capabilities": ["string"],
  "locations": ["string"],
  "issues": [
    {{
      "title": "string",
      "detail": "string",
      "penalty_weight": 0
    }}
  ]
}}

Profile filename: {filename or ""}
Profile text:
{text[:22000]}
""".strip()

    try:
        response = client.responses.create(model=OPENAI_MODEL, input=prompt)
        parsed = json_from_text(response.output_text)
        if parsed:
            parsed["_parse_mode"] = "openai"
        return parsed
    except Exception:
        return {}


def heuristic_profile_parse(text: str, filename: str | None = None) -> dict:
    lower = text.lower()

    industry = None
    rules = [
        ("Construction", ["construction", "contractor", "civil", "infrastructure", "building"]),
        ("ICT", ["software", "ict", "technology", "systems", "digital", "it services"]),
        ("Transport", ["transport", "shuttle", "fleet", "vehicle", "logistics"]),
        ("Professional Services", ["consulting", "advisory", "facilitation", "professional services"]),
        ("Tourism", ["tourism", "travel", "adventure", "hospitality", "destination"]),
        ("Security", ["security services", "guarding", "surveillance"]),
        ("Education", ["training provider", "education", "learnership", "skills development"]),
    ]
    for label, words in rules:
        if any(word in lower for word in words):
            industry = label
            break

    capabilities = []
    markers = ["services", "capabilities", "core services", "scope", "specialises in", "specializes in"]
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    for line in lines:
        ll = line.lower()
        if any(marker in ll for marker in markers):
            capabilities.extend(normalize_keywords(line))
    if not capabilities:
        capabilities = normalize_keywords(text[:3000])[:20]

    locations = []
    for province in [
        "Gauteng", "KwaZulu-Natal", "Western Cape", "Eastern Cape",
        "Free State", "Mpumalanga", "Limpopo", "North West", "Northern Cape",
    ]:
        if province.lower() in lower:
            locations.append(province)

    company_name = lines[0][:255] if lines else None
    if not company_name and filename:
        company_name = os.path.splitext(filename)[0]

    issues = []
    if not capabilities:
        issues.append({
            "title": "Capabilities are not clearly structured",
            "detail": "The profile does not clearly list service capabilities.",
            "penalty_weight": 6,
        })
    if not locations:
        issues.append({
            "title": "Operating locations are unclear",
            "detail": "The profile does not clearly indicate service provinces or locations.",
            "penalty_weight": 4,
        })
    if not industry:
        issues.append({
            "title": "Industry focus is unclear",
            "detail": "The profile does not strongly indicate the primary industry.",
            "penalty_weight": 5,
        })

    return {
        "company_name": company_name,
        "industry": industry,
        "capabilities": capabilities[:20],
        "locations": list(dict.fromkeys(locations)),
        "issues": issues,
        "_parse_mode": "heuristic",
    }


def parse_profile_text(text: str, filename: str | None = None) -> dict:
    parsed = parse_profile_with_openai(text, filename)
    if parsed:
        return parsed
    return heuristic_profile_parse(text, filename)


def serialize_profile(profile: Profile, include_issues: bool = True) -> dict:
    parsed = safe_loads(profile.parsed_json, {})
    data = {
        "id": profile.id,
        "name": profile.name,
        "company_name": profile.company_name,
        "original_filename": profile.original_filename,
        "industry": profile.industry,
        "capabilities_text": profile.capabilities_text,
        "locations_text": profile.locations_text,
        "extracted_text": profile.extracted_text,
        "parsed_json": parsed,
        "is_active": profile.is_active,
        "created_at": profile.created_at.isoformat() if profile.created_at else None,
        "updated_at": profile.updated_at.isoformat() if profile.updated_at else None,
    }
    if include_issues:
        issues = getattr(profile, "issues", []) or []
        data["issues"] = [
            {
                "id": issue.id,
                "issue_type": issue.issue_type,
                "title": issue.title,
                "detail": issue.detail,
                "status": issue.status,
                "penalty_weight": issue.penalty_weight,
            }
            for issue in issues
        ]
    return data


def tender_to_view_model(t: TenderCache) -> dict:
    return {
        "id": t.id,
        "tender_uid": t.tender_uid,
        "title": t.title,
        "description": t.description,
        "industry": t.industry,
        "tender_type": t.tender_type,
        "province": t.province,
        "buyer_name": t.buyer_name,
        "issued_date": t.issued_date.isoformat() if t.issued_date else None,
        "closing_date": t.closing_date.isoformat() if t.closing_date else None,
        "document_url": t.document_url,
        "source_url": t.source_url,
        "updated_at": t.updated_at.isoformat() if t.updated_at else None,
        "is_live": t.is_live,
    }


def build_profile_gap_summary(profile: Profile | None) -> dict:
    if not profile:
        return {"pending_count": 0, "fixed_count": 0, "penalty_total": 0.0}

    pending_count = 0
    fixed_count = 0
    penalty_total = 0.0
    for issue in (getattr(profile, "issues", []) or []):
        status = (getattr(issue, "status", "") or "").lower()
        penalty = float(getattr(issue, "penalty_weight", 0) or 0)
        if status == "pending":
            pending_count += 1
            penalty_total += penalty
        elif status == "fixed":
            fixed_count += 1

    return {
        "pending_count": pending_count,
        "fixed_count": fixed_count,
        "penalty_total": penalty_total,
    }


def keyword_overlap_score(profile: Profile | None, tender: TenderCache) -> float | None:
    if not profile:
        return None

    capabilities = [c.strip() for c in (profile.capabilities_text or "").split(",") if c.strip()]
    locations = [c.strip() for c in (profile.locations_text or "").split(",") if c.strip()]
    gap_summary = build_profile_gap_summary(profile)

    score = 22.0
    tender_blob = " ".join([
        tender.title or "",
        tender.description or "",
        tender.industry or "",
        tender.tender_type or "",
        tender.province or "",
        tender.buyer_name or "",
    ]).lower()

    if profile.industry and tender.industry and profile.industry.lower() == tender.industry.lower():
        score += 26.0
    elif profile.industry and profile.industry.lower() in tender_blob:
        score += 14.0

    matches = 0
    for capability in capabilities:
        if capability.lower() in tender_blob:
            matches += 1
    score += min(matches * 7.0, 35.0)

    for loc in locations:
        if tender.province and tender.province.lower() == loc.lower():
            score += 8.0
            break

    try:
        if tender.closing_date:
            days = (tender.closing_date - date.today()).days
            if days >= 0:
                score += 2.0 if days <= 7 else 5.0 if days <= 21 else 7.0
    except Exception:
        pass

    score -= min(gap_summary["penalty_total"], 18.0)
    score += min(gap_summary["fixed_count"] * 1.5, 4.0)

    return max(0.0, min(score, 100.0))


def build_fit_reasons(profile: Profile | None, tender: TenderCache) -> List[str]:
    if not profile:
        return []

    reasons = []
    capabilities = [x.strip() for x in (profile.capabilities_text or "").split(",") if x.strip()]
    locations = [x.strip() for x in (profile.locations_text or "").split(",") if x.strip()]
    tender_blob = f"{tender.title or ''} {tender.description or ''}".lower()

    matched_caps = [cap for cap in capabilities if cap.lower() in tender_blob]
    if matched_caps:
        reasons.append(f"Matches your service keywords: {', '.join(matched_caps[:4])}")

    if profile.industry and tender.industry and profile.industry.lower() == tender.industry.lower():
        reasons.append(f"Industry alignment with {tender.industry}")

    if locations and tender.province and any(loc.lower() == tender.province.lower() for loc in locations):
        reasons.append(f"Location alignment in {tender.province}")

    if tender.buyer_name:
        reasons.append(f"Issued by {tender.buyer_name}")

    if tender.closing_date:
        try:
            days = (tender.closing_date - date.today()).days
            if days >= 0:
                reasons.append(f"Still open with {days} day(s) remaining")
        except Exception:
            pass

    return reasons


def fit_band_from_score(score: Optional[float]) -> Optional[str]:
    if score is None:
        return None
    if score >= 80:
        return "high_potential"
    if score >= 55:
        return "possible_fit"
    return "low_fit"


def readiness_band_for_profile(profile: Profile | None) -> str:
    if not profile:
        return "no_profile"
    gap_summary = build_profile_gap_summary(profile)
    if gap_summary["pending_count"] == 0:
        return "ready"
    if gap_summary["pending_count"] <= 2:
        return "watchlist"
    return "needs_attention"


def readiness_message(profile: Profile | None, score: Optional[float]) -> str:
    if not profile:
        return "Upload and activate a profile to unlock TenderAI matching."
    gap_summary = build_profile_gap_summary(profile)
    if gap_summary["pending_count"] == 0:
        return "Your active profile looks ready for tender evaluation."
    if score is not None and score >= 70:
        return "This looks promising, but profile gaps may reduce readiness."
    return "Profile readiness gaps may weaken your bid position."


def build_fit_summary(score: Optional[float], reasons: List[str], profile: Profile | None = None) -> Optional[str]:
    if score is None:
        return None
    if reasons:
        base = "This tender could be a good fit because " + "; ".join(reasons[:3]) + "."
    elif score >= 60:
        base = "This tender shows promising metadata alignment with your active profile."
    else:
        base = "This tender has limited visible alignment from metadata alone and may require careful review."

    if profile:
        gap_summary = build_profile_gap_summary(profile)
        if gap_summary["pending_count"] > 0:
            base += f" Your active profile currently has {gap_summary['pending_count']} unresolved readiness gap(s)."
    return base


def build_document_match_check(tender: TenderCache, document_text: str) -> dict:
    text = (document_text or "").lower()
    if not text.strip():
        return {
            "document_match": False,
            "confidence": 0.0,
            "reason": "No extracted document text was available.",
            "matched_signals": [],
            "missing_signals": ["document text"],
        }

    matched_signals = []
    missing_signals = []
    score = 0.0

    title_tokens = tokenize(tender.title or "")[:8]
    title_matches = [tok for tok in title_tokens if tok in text]
    if len(title_matches) >= 2:
        matched_signals.append(f"Title overlap: {', '.join(title_matches[:5])}")
        score += min(len(title_matches) * 12.0, 36.0)
    else:
        missing_signals.append("strong title overlap")

    buyer_tokens = tokenize(tender.buyer_name or "")[:6]
    buyer_matches = [tok for tok in buyer_tokens if tok in text]
    if buyer_matches:
        matched_signals.append(f"Buyer overlap: {', '.join(buyer_matches[:4])}")
        score += min(len(buyer_matches) * 10.0, 25.0)
    elif tender.buyer_name:
        missing_signals.append("buyer name overlap")

    desc_tokens = tokenize(tender.description or "")[:10]
    desc_matches = [tok for tok in desc_tokens if tok in text]
    if len(desc_matches) >= 2:
        matched_signals.append(f"Description overlap: {', '.join(desc_matches[:5])}")
        score += min(len(desc_matches) * 4.0, 20.0)

    if tender.industry and tender.industry.lower() in text:
        matched_signals.append(f"Industry reference: {tender.industry}")
        score += 10.0

    if tender.province and tender.province.lower() in text:
        matched_signals.append(f"Province reference: {tender.province}")
        score += 6.0

    confidence = max(0.0, min(score, 100.0))
    is_match = confidence >= 35.0 and (len(title_matches) >= 2 or buyer_matches)

    return {
        "document_match": is_match,
        "confidence": confidence,
        "reason": "Document content appears consistent with the selected tender." if is_match else "Document content does not sufficiently match the selected tender metadata.",
        "matched_signals": matched_signals,
        "missing_signals": missing_signals,
    }


def extract_procurement_fields_fallback(tender: TenderCache, document_text: str) -> dict:
    text = document_text or ""
    briefing_date = None
    m = re.search(r"(brief(?:ing)?(?: session)?)[^0-9]{0,20}(\d{4}[-/]\d{2}[-/]\d{2}|\d{2}[-/]\d{2}[-/]\d{4})", text, re.I)
    if m:
        briefing_date = m.group(2).replace("/", "-")

    email_match = re.search(r"([A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,})", text)
    phone_match = re.search(r"(\+?\d[\d\s\-()]{7,}\d)", text)

    proposal_required = bool(re.search(r"\bproposal\b|\btechnical proposal\b|\bfinancial proposal\b", text, re.I))

    scope_summary = tender.description or "Scope to be confirmed from tender document."

    return {
        "briefing_date": briefing_date,
        "contact_person": None,
        "contact_email": email_match.group(1) if email_match else None,
        "contact_phone": phone_match.group(1) if phone_match else None,
        "proposal_required": proposal_required,
        "scope_summary": scope_summary[:700],
    }


def estimate_commercial_fallback(tender: TenderCache, profile: Profile, document_text: str) -> dict:
    text = (document_text or "").lower()
    title_blob = f"{tender.title or ''} {tender.description or ''}".lower()

    cost_drivers = []
    if any(w in text for w in ["delivery", "transport", "logistics", "site"]):
        cost_drivers.append("Delivery / logistics")
    if any(w in text for w in ["materials", "equipment", "supply"]):
        cost_drivers.append("Materials or equipment inputs")
    if any(w in text for w in ["staff", "labour", "training", "support"]):
        cost_drivers.append("Labour and delivery team capacity")
    if any(w in text for w in ["maintenance", "support", "service level", "sla"]):
        cost_drivers.append("Ongoing support or service obligations")
    if not cost_drivers:
        cost_drivers.append("Scope-specific execution costs need validation")

    pricing_complexity = "medium"
    if any(w in text for w in ["bill of quantities", "boq", "pricing schedule", "rate per", "unit price"]):
        pricing_complexity = "high"
    elif any(w in title_blob for w in ["appointment of panel", "framework", "supply and delivery"]):
        pricing_complexity = "medium"
    else:
        pricing_complexity = "low"

    potential_revenue_range = "Unknown — tender document does not clearly reveal contract value."
    if any(w in text for w in ["multi-year", "36 months", "three years", "framework", "panel"]):
        potential_revenue_range = "Potentially medium to high, depending on awarded scope and call-off volume."
    elif any(w in text for w in ["once-off", "one-time", "single project"]):
        potential_revenue_range = "Potentially low to medium, likely linked to a once-off scope."
    elif any(w in title_blob for w in ["supply", "services", "maintenance"]):
        potential_revenue_range = "Potentially medium, subject to actual quantities and pricing schedule."

    margin_risk = "medium"
    if pricing_complexity == "high":
        margin_risk = "high"
    elif pricing_complexity == "low":
        margin_risk = "low"

    return {
        "estimated_cost_drivers": cost_drivers,
        "pricing_complexity": pricing_complexity,
        "potential_revenue_range": potential_revenue_range,
        "margin_risk": margin_risk,
    }


def build_default_workspace() -> dict:
    return {
        "pursuit_status": "not_decided",
        "submission_readiness": "not_started",
        "next_action": "",
        "internal_notes": "",
        "document_override_status": "none",
        "checklist_status": "not_started",
    }


def merge_workspace(raw: dict) -> dict:
    workspace = build_default_workspace()
    workspace.update(raw.get("workspace") or {})
    return workspace


def build_checklist_fallback(analysis: dict) -> dict:
    mandatory = []
    compliance = []
    technical = []
    pricing = []

    if analysis.get("proposal_required"):
        mandatory.append("Prepare technical and commercial proposal pack")
        technical.append("Draft technical response aligned to scope")
        pricing.append("Prepare pricing schedule / commercial response")

    if analysis.get("briefing_date"):
        mandatory.append("Confirm briefing attendance requirements")

    compliance.append("Confirm company registration and statutory documents")
    compliance.append("Confirm tax / compliance documents before submission")
    technical.append("Review scope and delivery approach")
    pricing.append("Confirm costing assumptions and margin position")

    return {
        "mandatory_documents": mandatory,
        "compliance_items": compliance,
        "technical_items": technical,
        "pricing_items": pricing,
    }


def build_go_no_go(analysis: dict, profile: Optional[Profile]) -> dict:
    score = float(analysis.get("score") or 0)
    document_match = analysis.get("document_match")
    pricing_complexity = (analysis.get("pricing_complexity") or "medium").lower()
    margin_risk = (analysis.get("margin_risk") or "medium").lower()
    briefing_date = analysis.get("briefing_date")
    gap_summary = build_profile_gap_summary(profile)

    reasons = []
    decision = "pursue_with_caution"

    if document_match is False:
        decision = "do_not_pursue"
        reasons.append("Document match is not reliable.")
    elif score >= 75 and gap_summary["pending_count"] <= 2 and margin_risk != "high":
        decision = "pursue"
        reasons.append("Strong fit and manageable readiness risk.")
    elif score < 50:
        decision = "do_not_pursue"
        reasons.append("Low opportunity fit.")
    else:
        reasons.append("Opportunity has potential but needs controlled review.")

    if gap_summary["pending_count"] > 2:
        reasons.append(f"{gap_summary['pending_count']} readiness gap(s) remain unresolved.")

    if pricing_complexity == "high":
        reasons.append("Pricing complexity is high.")

    if margin_risk == "high":
        reasons.append("Margin risk is high.")

    if briefing_date:
        d = None
        try:
            d = (date.fromisoformat(briefing_date[:10]) - date.today()).days
        except Exception:
            d = None
        if d is not None and d <= 3:
            reasons.append("Briefing timeline is tight.")

    return {
        "decision": decision,
        "reasons": reasons,
    }


def build_executive_summary(tender: TenderCache, analysis: dict, workspace: dict, profile: Optional[Profile]) -> dict:
    go_no_go = build_go_no_go(analysis, profile)
    return {
        "title": tender.title or "Tender Opportunity",
        "department": tender.buyer_name,
        "closing_date": tender.closing_date.isoformat() if tender.closing_date else None,
        "fit_score": analysis.get("score"),
        "document_match": analysis.get("document_match"),
        "summary": analysis.get("summary"),
        "commercial_signal": analysis.get("potential_revenue_range"),
        "pricing_complexity": analysis.get("pricing_complexity"),
        "margin_risk": analysis.get("margin_risk"),
        "next_action": workspace.get("next_action"),
        "pursuit_status": workspace.get("pursuit_status"),
        "submission_readiness": workspace.get("submission_readiness"),
        "go_no_go": go_no_go,
    }


def latest_analysis_for(session, tender_id: int, profile_id: Optional[int]) -> Optional[dict]:
    if not profile_id:
        return None

    job = session.execute(
        select(AnalysisJob)
        .where(AnalysisJob.tender_id == tender_id, AnalysisJob.profile_id == profile_id)
        .order_by(desc(AnalysisJob.updated_at))
        .limit(1)
    ).scalars().first()

    if not job:
        return None

    raw = safe_loads(job.raw_result_json, {})
    workspace = merge_workspace(raw)
    checklist = raw.get("submission_checklist") or build_checklist_fallback(raw)
    go_no_go = raw.get("go_no_go") or build_go_no_go(raw, None)
    executive_summary = raw.get("executive_summary")

    return {
        "id": job.id,
        "status": job.status,
        "score": job.score,
        "summary": job.summary,
        "strengths": raw.get("strengths") or [x.strip() for x in (job.strengths_text or "").split("\n") if x.strip()],
        "gaps": raw.get("gaps") or [],
        "risks": raw.get("risks") or [x.strip() for x in (job.risks_text or "").split("\n") if x.strip()],
        "recommendation": job.recommendations_text,
        "error_message": job.error_message,
        "updated_at": job.updated_at.isoformat() if job.updated_at else None,
        "document_match": raw.get("document_match"),
        "document_match_confidence": raw.get("document_match_confidence"),
        "document_match_reason": raw.get("document_match_reason"),
        "matched_signals": raw.get("matched_signals") or [],
        "missing_signals": raw.get("missing_signals") or [],
        "estimated_cost_drivers": raw.get("estimated_cost_drivers") or [],
        "pricing_complexity": raw.get("pricing_complexity"),
        "potential_revenue_range": raw.get("potential_revenue_range"),
        "margin_risk": raw.get("margin_risk"),
        "briefing_date": raw.get("briefing_date"),
        "contact_person": raw.get("contact_person"),
        "contact_email": raw.get("contact_email"),
        "contact_phone": raw.get("contact_phone"),
        "proposal_required": raw.get("proposal_required"),
        "scope_summary": raw.get("scope_summary"),
        "workspace": workspace,
        "submission_checklist": checklist,
        "go_no_go": go_no_go,
        "executive_summary": executive_summary,
    }


def latest_analysis_map_for_tenders(session, profile_id: Optional[int], tender_ids: List[int]) -> Dict[int, dict]:
    if not profile_id or not tender_ids:
        return {}

    jobs = session.execute(
        select(AnalysisJob)
        .where(
            AnalysisJob.profile_id == profile_id,
            AnalysisJob.tender_id.in_(tender_ids),
        )
        .order_by(desc(AnalysisJob.updated_at))
    ).scalars().all()

    out: Dict[int, dict] = {}
    for job in jobs:
        if job.tender_id in out:
            continue
        raw = safe_loads(job.raw_result_json, {})
        out[job.tender_id] = {
            "score": job.score,
            "summary": job.summary,
            "briefing_date": raw.get("briefing_date"),
            "contact_person": raw.get("contact_person"),
            "contact_email": raw.get("contact_email"),
            "contact_phone": raw.get("contact_phone"),
            "proposal_required": raw.get("proposal_required"),
            "scope_summary": raw.get("scope_summary"),
            "document_match": raw.get("document_match"),
            "workspace": merge_workspace(raw),
            "status": job.status,
            "go_no_go": raw.get("go_no_go") or build_go_no_go(raw, None),
        }
    return out


def current_document_status(session, tender_id: int) -> Optional[dict]:
    doc = session.execute(
        select(TenderDocumentCache)
        .where(TenderDocumentCache.tender_id == tender_id)
        .order_by(desc(TenderDocumentCache.fetched_at), desc(TenderDocumentCache.id))
        .limit(1)
    ).scalars().first()
    if not doc:
        return None
    return {
        "fetch_status": doc.fetch_status,
        "error_message": doc.error_message,
        "filename": doc.filename,
        "content_type": doc.content_type,
        "fetched_at": doc.fetched_at.isoformat() if doc.fetched_at else None,
        "document_url": doc.document_url,
    }


def fetch_tender_document(session, tender: TenderCache, force_refresh: bool = False) -> dict:
    document_url = tender.document_url or tender.source_url
    if not document_url:
        return {"ok": False, "error": "No tender document URL available."}

    doc = session.execute(
        select(TenderDocumentCache)
        .where(TenderDocumentCache.tender_id == tender.id, TenderDocumentCache.document_url == document_url)
        .limit(1)
    ).scalars().first()

    if not doc:
        doc = TenderDocumentCache(tender_id=tender.id, document_url=document_url)
        session.add(doc)
        session.flush()

    if force_refresh:
        doc.fetch_status = "refreshing"
        doc.error_message = None
        session.flush()

    try:
        response = requests.get(document_url, timeout=25, allow_redirects=True)
        response.raise_for_status()

        content_type = (response.headers.get("Content-Type") or "").lower()
        filename = document_url.split("/")[-1][:255] if "/" in document_url else None
        lower_name = (filename or "").lower()
        lower_url = (document_url or "").lower()

        is_docx = (
            "wordprocessingml" in content_type
            or "msword" in content_type and lower_name.endswith(".docx")
            or lower_name.endswith(".docx")
            or lower_url.endswith(".docx")
        )

        is_doc = (
            content_type == "application/msword"
            or lower_name.endswith(".doc")
            or lower_url.endswith(".doc")
        ) and not is_docx

        is_pdf = (
            "pdf" in content_type
            or lower_name.endswith(".pdf")
            or lower_url.endswith(".pdf")
        )

        if is_doc:
            doc.filename = filename
            doc.content_type = content_type[:100] if content_type else None
            doc.binary_content = response.content
            doc.extracted_text = ""
            doc.fetch_status = "parse_failed"
            doc.error_message = "Legacy .doc files are not supported yet. Please convert to PDF or .docx."
            doc.fetched_at = utcnow()
            session.flush()
            return {"ok": False, "error": doc.error_message}

        if is_docx:
            suffix = ".docx"
        elif is_pdf:
            suffix = ".pdf"
        else:
            doc.filename = filename
            doc.content_type = content_type[:100] if content_type else None
            doc.binary_content = response.content
            doc.extracted_text = ""
            doc.fetch_status = "parse_failed"
            doc.error_message = f"Unsupported document type: {content_type or 'unknown'}"
            doc.fetched_at = utcnow()
            session.flush()
            return {"ok": False, "error": doc.error_message}

        extracted_text = ""
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(response.content)
            temp_path = tmp.name

        try:
            if suffix == ".pdf":
                extracted_text = extract_pdf_text(temp_path)
            elif suffix == ".docx":
                extracted_text = extract_docx_text(temp_path)
        finally:
            try:
                os.unlink(temp_path)
            except Exception:
                pass

        if not extracted_text.strip():
            doc.filename = filename
            doc.content_type = content_type[:100] if content_type else None
            doc.binary_content = response.content
            doc.extracted_text = ""
            doc.fetch_status = "parse_failed"
            doc.error_message = f"{suffix} was fetched but readable text could not be extracted."
            doc.fetched_at = utcnow()
            session.flush()
            return {"ok": False, "error": doc.error_message}

        doc.filename = filename
        doc.content_type = content_type[:100] if content_type else None
        doc.binary_content = response.content
        doc.extracted_text = extracted_text
        doc.fetch_status = "fetched"
        doc.error_message = None
        doc.fetched_at = utcnow()
        session.flush()

        return {"ok": True, "text_len": len(extracted_text), "content_type": content_type, "filename": filename}
    except Exception as exc:
        doc.fetch_status = "fetch_failed"
        doc.error_message = str(exc)
        session.flush()
        return {"ok": False, "error": str(exc)}


def analyze_tender_against_profile(tender: TenderCache, profile: Profile, document_text: str) -> dict:
    client = get_openai_client()

    profile_json = safe_loads(profile.parsed_json, {})
    tender_vm = tender_to_view_model(tender)
    match_check = build_document_match_check(tender, document_text)
    extracted_fields = extract_procurement_fields_fallback(tender, document_text)

    if not match_check["document_match"]:
        commercial = estimate_commercial_fallback(tender, profile, document_text)
        result = {
            "document_match": False,
            "document_match_confidence": match_check["confidence"],
            "document_match_reason": match_check["reason"],
            "matched_signals": match_check["matched_signals"],
            "missing_signals": match_check["missing_signals"],
            "score": 0,
            "summary": "TenderAI could not trust this analysis because the fetched document may belong to a different tender or a generic source page.",
            "strengths": [],
            "gaps": ["Document match could not be confirmed for this tender."],
            "risks": ["Analysis was blocked to avoid scoring the wrong tender document."],
            "recommendation": "Verify the tender document URL and retry analysis.",
            "estimated_cost_drivers": commercial["estimated_cost_drivers"],
            "pricing_complexity": commercial["pricing_complexity"],
            "potential_revenue_range": commercial["potential_revenue_range"],
            "margin_risk": commercial["margin_risk"],
            **extracted_fields,
            "_analysis_mode": "document_validation_failed",
        }
        result["workspace"] = build_default_workspace()
        result["submission_checklist"] = build_checklist_fallback(result)
        result["go_no_go"] = build_go_no_go(result, profile)
        result["executive_summary"] = build_executive_summary(tender, result, result["workspace"], profile)
        return result

    if client and document_text.strip():
        prompt = f"""
You are TenderAI, a procurement intelligence assistant.
First verify that the tender document belongs to the selected tender.
If it does not match, return document_match=false and do not produce a meaningful opportunity score.
If it matches, score and interpret the tender against the supplier profile.

Also extract procurement fields where possible:
- briefing_date
- contact_person
- contact_email
- contact_phone
- proposal_required
- scope_summary

Also estimate indicative commercial intelligence:
- estimated_cost_drivers
- pricing_complexity
- potential_revenue_range
- margin_risk

Also build a submission checklist:
- mandatory_documents
- compliance_items
- technical_items
- pricing_items

Also return a final go/no-go block:
- decision (pursue|pursue_with_caution|do_not_pursue)
- reasons

Return JSON only.

Schema:
{{
  "document_match": true,
  "document_match_reason": "string",
  "score": 0,
  "summary": "string",
  "strengths": ["string"],
  "gaps": ["string"],
  "risks": ["string"],
  "recommendation": "string",
  "estimated_cost_drivers": ["string"],
  "pricing_complexity": "low|medium|high",
  "potential_revenue_range": "string",
  "margin_risk": "low|medium|high",
  "briefing_date": "string or null",
  "contact_person": "string or null",
  "contact_email": "string or null",
  "contact_phone": "string or null",
  "proposal_required": true,
  "scope_summary": "string",
  "submission_checklist": {{
    "mandatory_documents": ["string"],
    "compliance_items": ["string"],
    "technical_items": ["string"],
    "pricing_items": ["string"]
  }},
  "go_no_go": {{
    "decision": "pursue",
    "reasons": ["string"]
  }}
}}

Supplier profile JSON:
{json.dumps(profile_json, ensure_ascii=False, default=str)[:12000]}

Tender metadata:
{json.dumps(tender_vm, ensure_ascii=False, default=str)[:6000]}

Pre-check:
{json.dumps(match_check, ensure_ascii=False, default=str)[:2000]}

Tender document text:
{document_text[:22000]}
""".strip()

        try:
            response = client.responses.create(model=OPENAI_MODEL, input=prompt)
            parsed = json_from_text(response.output_text)
            if parsed:
                parsed["_analysis_mode"] = "openai"
                parsed.setdefault("document_match", True)
                parsed.setdefault("document_match_confidence", match_check["confidence"])
                parsed.setdefault("document_match_reason", match_check["reason"])
                parsed.setdefault("matched_signals", match_check["matched_signals"])
                parsed.setdefault("missing_signals", match_check["missing_signals"])
                parsed.setdefault("briefing_date", extracted_fields["briefing_date"])
                parsed.setdefault("contact_person", extracted_fields["contact_person"])
                parsed.setdefault("contact_email", extracted_fields["contact_email"])
                parsed.setdefault("contact_phone", extracted_fields["contact_phone"])
                parsed.setdefault("proposal_required", extracted_fields["proposal_required"])
                parsed.setdefault("scope_summary", extracted_fields["scope_summary"])
                parsed.setdefault("workspace", build_default_workspace())
                parsed.setdefault("submission_checklist", build_checklist_fallback(parsed))
                parsed.setdefault("go_no_go", build_go_no_go(parsed, profile))
                parsed.setdefault("executive_summary", build_executive_summary(tender, parsed, parsed["workspace"], profile))

                if parsed.get("document_match") is False:
                    parsed["score"] = 0
                    parsed.setdefault("recommendation", "Verify the tender document URL and retry analysis.")
                return parsed
        except Exception:
            pass

    score = keyword_overlap_score(profile, tender) or 0
    reasons = build_fit_reasons(profile, tender)
    commercial = estimate_commercial_fallback(tender, profile, document_text)

    result = {
        "document_match": True,
        "document_match_confidence": match_check["confidence"],
        "document_match_reason": match_check["reason"],
        "matched_signals": match_check["matched_signals"],
        "missing_signals": match_check["missing_signals"],
        "score": score,
        "summary": build_fit_summary(score, reasons, profile) or "Fallback analysis was used.",
        "strengths": reasons[:4],
        "gaps": ["Detailed AI interpretation was unavailable."],
        "risks": [] if score >= 60 else ["Profile alignment appears limited from the available metadata."],
        "recommendation": "Proceed to full response preparation." if score >= 60 else "Review carefully before committing resources.",
        "estimated_cost_drivers": commercial["estimated_cost_drivers"],
        "pricing_complexity": commercial["pricing_complexity"],
        "potential_revenue_range": commercial["potential_revenue_range"],
        "margin_risk": commercial["margin_risk"],
        **extracted_fields,
        "_analysis_mode": "heuristic",
        "workspace": build_default_workspace(),
    }
    result["submission_checklist"] = build_checklist_fallback(result)
    result["go_no_go"] = build_go_no_go(result, profile)
    result["executive_summary"] = build_executive_summary(tender, result, result["workspace"], profile)
    return result


def generate_proposal_content(tender: TenderCache, profile: Profile, analysis: dict) -> dict:
    client = get_openai_client()

    profile_json = safe_loads(profile.parsed_json, {})
    tender_vm = tender_to_view_model(tender)

    if client:
        prompt = f"""
You are TenderAI. Draft a procurement proposal response for the supplier.
Write a professional proposal structure suitable for submission or adaptation.
Return JSON only.

Schema:
{{
  "title": "string",
  "executive_summary": "string",
  "company_positioning": "string",
  "approach": "string",
  "compliance_notes": ["string"],
  "deliverables": ["string"],
  "commercial_notes": "string",
  "closing_statement": "string"
}}

Supplier profile JSON:
{json.dumps(profile_json, ensure_ascii=False, default=str)[:12000]}

Tender metadata:
{json.dumps(tender_vm, ensure_ascii=False, default=str)[:7000]}

Tender analysis JSON:
{json.dumps(analysis, ensure_ascii=False, default=str)[:10000]}
""".strip()

        try:
            response = client.responses.create(model=OPENAI_MODEL, input=prompt)
            parsed = json_from_text(response.output_text)
            if parsed:
                parsed["_proposal_mode"] = "openai"
                return parsed
        except Exception:
            pass

    company_name = profile.company_name or profile.name or "The Supplier"
    return {
        "title": f"Proposal Response: {tender.title or 'Tender Opportunity'}",
        "executive_summary": f"{company_name} submits this response in relation to the opportunity issued by {tender.buyer_name or 'the department'}.",
        "company_positioning": f"{company_name} is positioned to deliver within the required scope, supported by relevant capabilities captured in the active supplier profile.",
        "approach": "Our team will review the final tender requirements, align delivery resources, confirm compliance items, and structure execution around the tender scope and timelines.",
        "compliance_notes": analysis.get("gaps") or ["Final compliance review is required before submission."],
        "deliverables": analysis.get("strengths") or ["Delivery aligned to tender scope and requirements."],
        "commercial_notes": f"Indicative commercial view: {analysis.get('potential_revenue_range', 'Value to be confirmed')} with margin risk assessed as {analysis.get('margin_risk', 'medium')}.",
        "closing_statement": "We welcome the opportunity to submit a compliant and competitive response and remain available for clarification or presentation.",
        "_proposal_mode": "heuristic",
    }


def build_proposal_docx(proposal: dict, tender: TenderCache, profile: Profile) -> bytes:
    if DocxDocument is None:
        raise RuntimeError("python-docx is not installed.")

    doc = DocxDocument()
    doc.add_heading(proposal.get("title") or f"Proposal: {tender.title}", 0)

    doc.add_paragraph(f"Tender: {tender.title or 'N/A'}")
    doc.add_paragraph(f"Department / Buyer: {tender.buyer_name or 'N/A'}")
    doc.add_paragraph(f"Supplier: {profile.company_name or profile.name or 'N/A'}")
    doc.add_paragraph(f"Closing Date: {tender.closing_date.isoformat() if tender.closing_date else 'N/A'}")

    sections = [
        ("Executive Summary", proposal.get("executive_summary")),
        ("Company Positioning", proposal.get("company_positioning")),
        ("Approach", proposal.get("approach")),
        ("Commercial Notes", proposal.get("commercial_notes")),
        ("Closing Statement", proposal.get("closing_statement")),
    ]

    for heading, body in sections:
        if body:
            doc.add_heading(heading, level=1)
            doc.add_paragraph(body)

    compliance_notes = proposal.get("compliance_notes") or []
    if compliance_notes:
        doc.add_heading("Compliance Notes", level=1)
        for item in compliance_notes:
            doc.add_paragraph(str(item), style="List Bullet")

    deliverables = proposal.get("deliverables") or []
    if deliverables:
        doc.add_heading("Deliverables / Strengths", level=1)
        for item in deliverables:
            doc.add_paragraph(str(item), style="List Bullet")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def find_running_analysis(session, tender_id: int, profile_id: int) -> Optional[AnalysisJob]:
    return session.execute(
        select(AnalysisJob)
        .where(
            AnalysisJob.tender_id == tender_id,
            AnalysisJob.profile_id == profile_id,
            AnalysisJob.status == "running",
        )
        .order_by(desc(AnalysisJob.updated_at))
        .limit(1)
    ).scalars().first()


@app.context_processor
def inject_globals():
    with get_db_session() as session:
        active_profile = get_active_profile(session)
        latest_ingest = session.execute(
            select(IngestRun).order_by(desc(IngestRun.started_at), desc(IngestRun.id)).limit(1)
        ).scalars().first()

        return {
            "active_profile": serialize_profile(active_profile, include_issues=False) if active_profile else None,
            "latest_ingest": {
                "status": latest_ingest.status,
                "started_at": latest_ingest.started_at.isoformat() if latest_ingest.started_at else None,
            } if latest_ingest else None,
            "today": date.today().isoformat(),
        }


@app.template_filter("days_left")
def days_left_filter(target_date):
    if not target_date:
        return None
    if isinstance(target_date, str):
        try:
            target_date = date.fromisoformat(target_date[:10])
        except Exception:
            return None
    return (target_date - date.today()).days


@app.get("/health")
def health():
    with get_db_session() as session:
        count = session.execute(select(func.count()).select_from(TenderCache)).scalar_one()
        return jsonify({"ok": True, "cached_tenders": count, "time": utcnow().isoformat()})


@app.get("/")
def home():
    with get_db_session() as session:
        active_profile = get_active_profile(session)
        readiness_band = readiness_band_for_profile(active_profile)
        readiness_note = readiness_message(active_profile, None)
        gap_summary = build_profile_gap_summary(active_profile)

        total_live = session.execute(
            select(func.count()).select_from(TenderCache).where(TenderCache.is_live.is_(True))
        ).scalar_one()

        tenders = session.execute(
            select(TenderCache)
            .where(TenderCache.is_live.is_(True))
            .order_by(TenderCache.closing_date.asc().nulls_last(), desc(TenderCache.updated_at))
            .limit(24)
        ).scalars().all()

        analysis_map = latest_analysis_map_for_tenders(
            session,
            active_profile.id if active_profile else None,
            [t.id for t in tenders],
        )

        featured = []
        for t in tenders:
            score = keyword_overlap_score(active_profile, t)
            reasons = build_fit_reasons(active_profile, t)
            fit_band = fit_band_from_score(score)
            tender_vm = tender_to_view_model(t)
            latest = analysis_map.get(t.id, {})

            featured.append({
                "tender": tender_vm,
                "score": score,
                "fit_band": fit_band,
                "fit_summary": latest.get("scope_summary") or build_fit_summary(score, reasons, active_profile),
                "fit_reasons": reasons,
                "readiness_band": readiness_band,
                "briefing_date": latest.get("briefing_date"),
                "contact_person": latest.get("contact_person"),
                "contact_email": latest.get("contact_email"),
                "contact_phone": latest.get("contact_phone"),
            })

        if active_profile:
            featured.sort(key=lambda x: (x["score"] or 0), reverse=True)

        return render_template(
            "home.html",
            total_live=total_live,
            featured=featured[:12],
            readiness_band=readiness_band,
            readiness_note=readiness_note,
            profile_gap_summary=gap_summary,
        )


@app.get("/tenders")
def tenders():
    with get_db_session() as session:
        active_profile = get_active_profile(session)
        readiness_band = readiness_band_for_profile(active_profile)
        gap_summary = build_profile_gap_summary(active_profile)

        province = (request.args.get("province") or "").strip()
        tender_type = (request.args.get("tender_type") or "").strip()
        industry = (request.args.get("industry") or "").strip()
        issued_from = (request.args.get("issued_from") or "").strip()
        search_text = (request.args.get("q") or "").strip()
        fit_band_filter = (request.args.get("fit_band") or "").strip()

        query = select(TenderCache).where(TenderCache.is_live.is_(True))

        if province:
            query = query.where(TenderCache.province == province)
        if tender_type:
            query = query.where(TenderCache.tender_type == tender_type)
        if industry:
            query = query.where(TenderCache.industry == industry)
        if issued_from:
            try:
                dt = datetime.strptime(issued_from, "%Y-%m-%d").date()
                query = query.where(TenderCache.issued_date >= dt)
            except ValueError:
                pass
        if search_text:
            like_term = f"%{search_text.lower()}%"
            query = query.where(
                or_(
                    func.lower(TenderCache.title).like(like_term),
                    func.lower(func.coalesce(TenderCache.description, "")).like(like_term),
                    func.lower(func.coalesce(TenderCache.buyer_name, "")).like(like_term),
                    func.lower(func.coalesce(TenderCache.industry, "")).like(like_term),
                )
            )

        items = session.execute(
            query.order_by(TenderCache.closing_date.asc().nulls_last(), desc(TenderCache.updated_at)).limit(200)
        ).scalars().all()

        analysis_map = latest_analysis_map_for_tenders(
            session,
            active_profile.id if active_profile else None,
            [t.id for t in items],
        )

        ranked = []
        for t in items:
            score = keyword_overlap_score(active_profile, t)
            reasons = build_fit_reasons(active_profile, t)
            band = fit_band_from_score(score)
            if fit_band_filter and band != fit_band_filter:
                continue

            latest = analysis_map.get(t.id, {})
            ranked.append({
                "tender": tender_to_view_model(t),
                "score": score,
                "fit_band": band,
                "fit_summary": latest.get("scope_summary") or build_fit_summary(score, reasons, active_profile),
                "fit_reasons": reasons,
                "readiness_band": readiness_band,
                "briefing_date": latest.get("briefing_date"),
                "contact_person": latest.get("contact_person"),
                "contact_email": latest.get("contact_email"),
                "contact_phone": latest.get("contact_phone"),
                "proposal_required": latest.get("proposal_required"),
                "analysis_status": latest.get("status"),
                "workspace": latest.get("workspace") or build_default_workspace(),
                "go_no_go": latest.get("go_no_go"),
            })

        if active_profile:
            ranked.sort(key=lambda x: (x["score"] or 0), reverse=True)

        provinces = session.execute(
            select(TenderCache.province).where(TenderCache.is_live.is_(True), TenderCache.province.is_not(None)).distinct().order_by(TenderCache.province)
        ).scalars().all()

        tender_types = session.execute(
            select(TenderCache.tender_type).where(TenderCache.is_live.is_(True), TenderCache.tender_type.is_not(None)).distinct().order_by(TenderCache.tender_type)
        ).scalars().all()

        industries = session.execute(
            select(TenderCache.industry).where(TenderCache.is_live.is_(True), TenderCache.industry.is_not(None)).distinct().order_by(TenderCache.industry)
        ).scalars().all()

        band_counts = {"high_potential": 0, "possible_fit": 0, "low_fit": 0}
        for item in ranked:
            if item["fit_band"] in band_counts:
                band_counts[item["fit_band"]] += 1

        return render_template(
            "feed.html",
            ranked_tenders=ranked,
            provinces=provinces,
            tender_types=tender_types,
            industries=industries,
            filters={
                "province": province,
                "tender_type": tender_type,
                "industry": industry,
                "issued_from": issued_from,
                "q": search_text,
                "fit_band": fit_band_filter,
            },
            readiness_band=readiness_band,
            profile_gap_summary=gap_summary,
            band_counts=band_counts,
        )


@app.get("/tender/<int:tender_id>")
def tender_detail(tender_id: int):
    with get_db_session() as session:
        tender = session.get(TenderCache, tender_id)
        if not tender:
            abort(404)

        active_profile = get_active_profile(session)
        score = keyword_overlap_score(active_profile, tender)
        reasons = build_fit_reasons(active_profile, tender)
        fit_band = fit_band_from_score(score)
        fit_summary = build_fit_summary(score, reasons, active_profile)
        latest_analysis = latest_analysis_for(session, tender_id, active_profile.id if active_profile else None)
        readiness_band = readiness_band_for_profile(active_profile)
        readiness_note = readiness_message(active_profile, score)
        gap_summary = build_profile_gap_summary(active_profile)
        doc_status = current_document_status(session, tender_id)
        running_job = None
        if active_profile:
            running_job = find_running_analysis(session, tender_id, active_profile.id)

        proposal_preview = None
        if latest_analysis and latest_analysis.get("proposal_required"):
            try:
                proposal_preview = generate_proposal_content(tender, active_profile, latest_analysis)
            except Exception:
                proposal_preview = None

        return render_template(
            "tender_detail.html",
            tender=tender_to_view_model(tender),
            alignment_score=score,
            fit_summary=fit_summary,
            fit_reasons=reasons,
            fit_band=fit_band,
            can_analyze=bool(active_profile),
            analyze_action_url=url_for("analyze_tender_page", tender_id=tender_id),
            latest_analysis=latest_analysis,
            readiness_band=readiness_band,
            readiness_note=readiness_note,
            profile_gap_summary=gap_summary,
            document_status=doc_status,
            running_job=running_job,
            proposal_preview=proposal_preview,
        )


@app.post("/tender/<int:tender_id>/analyze")
def analyze_tender_page(tender_id: int):
    with get_db_session() as session:
        tender = session.get(TenderCache, tender_id)
        if not tender:
            flash("Tender not found.", "error")
            return redirect(url_for("tenders"))

        active_profile = get_active_profile(session)
        if not active_profile:
            flash("Please upload and activate a business profile first.", "error")
            return redirect(url_for("profiles"))

        existing_running = find_running_analysis(session, tender_id, active_profile.id)
        if existing_running:
            flash("An analysis is already running for this tender and profile.", "warning")
            return redirect(url_for("tender_detail", tender_id=tender_id))

        job = AnalysisJob(profile_id=active_profile.id, tender_id=tender_id, status="running")
        session.add(job)
        session.flush()

        try:
            document_url = tender.document_url or tender.source_url
            if not document_url:
                job.status = "failed"
                job.error_message = "No tender document URL is available for this tender."
                flash(job.error_message, "error")
                return redirect(url_for("tender_detail", tender_id=tender_id))

            doc = session.execute(
                select(TenderDocumentCache)
                .where(TenderDocumentCache.tender_id == tender_id, TenderDocumentCache.document_url == document_url)
                .limit(1)
            ).scalars().first()

            if not doc or doc.fetch_status != "fetched":
                fetch_result = fetch_tender_document(session, tender)
                if not fetch_result.get("ok"):
                    job.status = "failed"
                    job.error_message = fetch_result.get("error") or "Unable to fetch tender document."
                    flash(f"Unable to fetch tender document: {job.error_message}", "error")
                    return redirect(url_for("tender_detail", tender_id=tender_id))

                doc = session.execute(
                    select(TenderDocumentCache)
                    .where(TenderDocumentCache.tender_id == tender_id, TenderDocumentCache.document_url == document_url)
                    .limit(1)
                ).scalars().first()

            extracted_text = (doc.extracted_text or "").strip() if doc else ""
            if not extracted_text:
                job.status = "failed"
                job.error_message = "Tender document was fetched, but no readable text was extracted."
                flash(job.error_message, "error")
                return redirect(url_for("tender_detail", tender_id=tender_id))

            analysis = analyze_tender_against_profile(tender, active_profile, extracted_text) or {}
            analysis.setdefault("workspace", build_default_workspace())
            analysis.setdefault("submission_checklist", build_checklist_fallback(analysis))
            analysis.setdefault("go_no_go", build_go_no_go(analysis, active_profile))
            analysis.setdefault("executive_summary", build_executive_summary(tender, analysis, analysis["workspace"], active_profile))

            job.status = "completed"
            job.score = float(analysis.get("score") or 0)
            job.summary = analysis.get("summary")
            job.strengths_text = "\n".join(analysis.get("strengths") or [])
            job.risks_text = "\n".join(analysis.get("risks") or [])
            job.recommendations_text = analysis.get("recommendation")
            job.raw_result_json = json.dumps(analysis, ensure_ascii=False, default=str)
            job.error_message = None

            doc.parsed_json = json.dumps({"analysis": analysis}, ensure_ascii=False, default=str)
            session.flush()

            if analysis.get("document_match") is False:
                flash("Tender analysis blocked because the fetched document may belong to another tender.", "warning")
            else:
                flash("Tender analysis completed.", "success")

            return redirect(url_for("tender_detail", tender_id=tender_id))

        except Exception as exc:
            job.status = "failed"
            job.error_message = str(exc)
            session.flush()
            flash(f"Analysis failed: {exc}", "error")
            return redirect(url_for("tender_detail", tender_id=tender_id))


@app.post("/tender/<int:tender_id>/workspace")
def update_workspace(tender_id: int):
    with get_db_session() as session:
        active_profile = get_active_profile(session)
        if not active_profile:
            flash("Please upload and activate a profile first.", "error")
            return redirect(url_for("profiles"))

        job = session.execute(
            select(AnalysisJob)
            .where(AnalysisJob.tender_id == tender_id, AnalysisJob.profile_id == active_profile.id)
            .order_by(desc(AnalysisJob.updated_at))
            .limit(1)
        ).scalars().first()

        if not job:
            flash("Analyze the tender first before updating the workspace.", "error")
            return redirect(url_for("tender_detail", tender_id=tender_id))

        raw = safe_loads(job.raw_result_json, {})
        workspace = merge_workspace(raw)
        workspace["pursuit_status"] = (request.form.get("pursuit_status") or workspace["pursuit_status"]).strip()
        workspace["submission_readiness"] = (request.form.get("submission_readiness") or workspace["submission_readiness"]).strip()
        workspace["next_action"] = (request.form.get("next_action") or "").strip()
        workspace["internal_notes"] = (request.form.get("internal_notes") or "").strip()
        workspace["document_override_status"] = (request.form.get("document_override_status") or workspace["document_override_status"]).strip()
        workspace["checklist_status"] = (request.form.get("checklist_status") or workspace["checklist_status"]).strip()

        raw["workspace"] = workspace
        raw["go_no_go"] = build_go_no_go(raw, active_profile)
        tender = session.get(TenderCache, tender_id)
        if tender:
            raw["executive_summary"] = build_executive_summary(tender, raw, workspace, active_profile)
        job.raw_result_json = json.dumps(raw, ensure_ascii=False, default=str)
        flash("Bid workspace updated.", "success")
        return redirect(url_for("tender_detail", tender_id=tender_id))


@app.post("/tender/<int:tender_id>/refetch-document")
def refetch_document_page(tender_id: int):
    with get_db_session() as session:
        tender = session.get(TenderCache, tender_id)
        if not tender:
            flash("Tender not found.", "error")
            return redirect(url_for("tenders"))

        result = fetch_tender_document(session, tender, force_refresh=True)
        if result.get("ok"):
            flash("Tender document re-fetched successfully.", "success")
        else:
            flash(f"Document re-fetch failed: {result.get('error')}", "error")
        return redirect(url_for("tender_detail", tender_id=tender_id))


@app.post("/tender/<int:tender_id>/reanalyze")
def reanalyze_tender_page(tender_id: int):
    return analyze_tender_page(tender_id)


@app.get("/api/tenders/<int:tender_id>/proposal")
def api_generate_proposal(tender_id: int):
    with get_db_session() as session:
        tender = session.get(TenderCache, tender_id)
        if not tender:
            return jsonify({"ok": False, "error": "Tender not found."}), 404

        active_profile = get_active_profile(session)
        if not active_profile:
            return jsonify({"ok": False, "error": "No active profile found."}), 400

        latest = latest_analysis_for(session, tender_id, active_profile.id)
        if not latest:
            return jsonify({"ok": False, "error": "No analysis found. Analyze the tender first."}), 400

        proposal = generate_proposal_content(tender, active_profile, latest)
        return jsonify({"ok": True, "proposal": proposal})


@app.get("/tender/<int:tender_id>/proposal.docx")
def download_proposal_docx(tender_id: int):
    with get_db_session() as session:
        tender = session.get(TenderCache, tender_id)
        if not tender:
            abort(404)

        active_profile = get_active_profile(session)
        if not active_profile:
            flash("No active profile found.", "error")
            return redirect(url_for("profiles"))

        latest = latest_analysis_for(session, tender_id, active_profile.id)
        if not latest:
            flash("Analyze the tender before generating a proposal.", "error")
            return redirect(url_for("tender_detail", tender_id=tender_id))

        proposal = generate_proposal_content(tender, active_profile, latest)

        try:
            payload = build_proposal_docx(proposal, tender, active_profile)
        except Exception as exc:
            flash(f"Unable to generate DOCX proposal: {exc}", "error")
            return redirect(url_for("tender_detail", tender_id=tender_id))

        return send_file(
            io.BytesIO(payload),
            as_attachment=True,
            download_name=f"proposal_tender_{tender_id}.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )


@app.get("/profiles")
def profiles():
    with get_db_session() as session:
        profiles_list = session.execute(select(Profile).order_by(desc(Profile.updated_at))).scalars().all()
        return render_template("profiles.html", profiles=[serialize_profile(p, include_issues=True) for p in profiles_list])


@app.post("/profiles")
@app.post("/profiles/upload")
def upload_profile():
    uploaded = request.files.get("profile_pdf") or request.files.get("profile_file") or request.files.get("file")
    if not uploaded or not uploaded.filename.lower().endswith(".pdf"):
        flash("Please upload a PDF profile.", "error")
        return redirect(url_for("profiles"))

    try:
        text = extract_pdf_text(uploaded)
    except Exception as exc:
        flash(f"Could not read PDF: {exc}", "error")
        return redirect(url_for("profiles"))

    parsed = parse_profile_text(text, uploaded.filename)

    with get_db_session() as session:
        session.execute(Profile.__table__.update().values(is_active=False))
        profile = Profile(
            name=parsed.get("company_name") or os.path.splitext(uploaded.filename)[0],
            company_name=parsed.get("company_name"),
            original_filename=uploaded.filename,
            industry=parsed.get("industry"),
            capabilities_text=", ".join(parsed.get("capabilities") or []),
            locations_text=", ".join(parsed.get("locations") or []),
            extracted_text=text[:200000],
            parsed_json=json.dumps(parsed, ensure_ascii=False, default=str),
            is_active=True,
        )
        session.add(profile)
        session.flush()
        upsert_profile_issues(session, profile, parsed.get("issues") or [])

    flash(f"Profile uploaded and set as active. Parse mode: {parsed.get('_parse_mode', 'heuristic')}.", "success")
    return redirect(url_for("profiles"))


@app.post("/profiles/<int:profile_id>/activate")
def activate_profile(profile_id: int):
    with get_db_session() as session:
        profile = session.get(Profile, profile_id)
        if not profile:
            flash("Profile not found.", "error")
            return redirect(url_for("profiles"))

        session.execute(Profile.__table__.update().values(is_active=False))
        profile.is_active = True
        flash("Active profile updated.", "success")
        return redirect(url_for("profiles"))


@app.post("/profile-issues/<int:issue_id>/status")
def update_issue_status(issue_id: int):
    status = (request.form.get("status") or "").strip().lower()
    if status not in {"pending", "fixed"}:
        flash("Invalid issue status.", "error")
        return redirect(url_for("profiles"))

    with get_db_session() as session:
        issue = session.get(ProfileIssue, issue_id)
        if not issue:
            flash("Issue not found.", "error")
            return redirect(url_for("profiles"))
        issue.status = status
        flash("Issue status updated.", "success")
        return redirect(url_for("profiles"))


@app.get("/api/admin/openai-status")
def admin_openai_status():
    if not admin_allowed():
        return jsonify({"ok": False, "error": "unauthorized"}), 403
    client = get_openai_client()
    return jsonify({
        "ok": True,
        "api_key_present": bool((os.getenv("OPENAI_API_KEY") or "").strip()),
        "client_ready": client is not None,
        "model": OPENAI_MODEL,
    })


@app.post("/api/admin/run-ingest")
@app.get("/api/admin/run-ingest")
def api_run_ingest():
    if not admin_allowed():
        return jsonify({"ok": False, "error": "unauthorized"}), 403

    with get_db_session() as session:
        result = ingest_tenders(session=session, max_pages=int(os.getenv("INGEST_MAX_PAGES", "1")))
        return jsonify(result)


@app.post("/api/admin/fetch-documents")
@app.get("/api/admin/fetch-documents")
def api_fetch_documents():
    if not admin_allowed():
        return jsonify({"ok": False, "error": "unauthorized"}), 403

    limit = max(1, min(int(request.args.get("limit", 20)), 100))
    fetched = 0
    errors = []

    with get_db_session() as session:
        tenders = session.execute(select(TenderCache).order_by(desc(TenderCache.updated_at)).limit(limit)).scalars().all()

        for tender in tenders:
            result = fetch_tender_document(session, tender)
            if result.get("ok"):
                fetched += 1
            else:
                errors.append({"tender_id": tender.id, "error": result.get("error")})

    return jsonify({"ok": True, "requested": limit, "fetched": fetched, "errors": errors})


@app.errorhandler(404)
def not_found(_):
    return render_template("404.html"), 404


@app.errorhandler(500)
def internal_error(_):
    return render_template("500.html"), 500


if __name__ == "__main__":
    port = int(os.getenv("PORT", "8080"))
    app.run(host="0.0.0.0", port=port, debug=False)
